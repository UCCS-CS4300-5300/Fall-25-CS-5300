"""
Tests for Resume Parsing Feature (Issues #48, #49, #50)

Tests cover:
- openai_utils.py: OpenAI client initialization and availability
- resume_parser.py: AI-powered resume parsing
- views.py: Integration of parsing into upload flow
"""

import json
from io import BytesIO
from unittest.mock import patch, MagicMock
from django.contrib.auth.models import User
from django.test import TestCase, override_settings
from django.urls import reverse
from django.core.files.uploadedfile import SimpleUploadedFile
from docx import Document as DocxDocument

from ..models import UploadedResume
from ..openai_utils import get_openai_client, ai_available, MAX_TOKENS
from ..resume_parser import parse_resume_with_ai, validate_parsed_data
from .test_credentials import TEST_PASSWORD


# === Helper Functions ===

def create_test_user():
    """Create a test user for authentication."""
    return User.objects.create_user(
        username="testuser",
        password=TEST_PASSWORD,
        email="test@example.com"
    )


def create_sample_resume_text():
    """Generate sample resume content for testing."""
    return """
    John Doe
    Software Engineer

    Skills: Python, Django, JavaScript, React, SQL, Git

    Experience:
    Senior Developer at Tech Corp (2020-2023)
    - Led development of web applications
    - Managed team of 5 developers

    Junior Developer at StartUp Inc (2018-2020)
    - Built REST APIs with Django
    - Implemented frontend features

    Education:
    Bachelor of Science in Computer Science
    State University, 2018
    """


def create_sample_parsed_data():
    """Generate sample parsed data structure."""
    return {
        "skills": ["Python", "Django", "JavaScript", "React", "SQL", "Git"],
        "experience": [
            {
                "company": "Tech Corp",
                "title": "Senior Developer",
                "duration": "2020-2023",
                "description": "Led development of web applications"
            },
            {
                "company": "StartUp Inc",
                "title": "Junior Developer",
                "duration": "2018-2020",
                "description": "Built REST APIs with Django"
            }
        ],
        "education": [
            {
                "institution": "State University",
                "degree": "Bachelor of Science",
                "field": "Computer Science",
                "year": "2018"
            }
        ]
    }


# === Tests for openai_utils.py ===

class OpenAIUtilsTests(TestCase):
    """Test OpenAI utility functions."""

    @override_settings(OPENAI_API_KEY="test-key-12345")
    @patch('active_interview_app.openai_utils.OpenAI')
    def test_get_openai_client_success(self, mock_openai):
        """Test successful OpenAI client initialization."""
        # Reset the global client
        import active_interview_app.openai_utils as utils
        utils._openai_client = None

        # Mock OpenAI client
        mock_client = MagicMock()
        mock_openai.return_value = mock_client

        # Get client
        client = get_openai_client()

        # Verify
        self.assertIsNotNone(client)
        mock_openai.assert_called_once_with(api_key="test-key-12345")

    @override_settings(OPENAI_API_KEY="")
    def test_get_openai_client_no_api_key(self):
        """Test OpenAI client initialization fails without API key."""
        # Reset the global client
        import active_interview_app.openai_utils as utils
        utils._openai_client = None

        # Attempt to get client without API key
        with self.assertRaises(ValueError) as context:
            get_openai_client()

        self.assertIn("No OpenAI API key available", str(context.exception))

    @override_settings(OPENAI_API_KEY="test-key")
    @patch('active_interview_app.openai_utils.OpenAI')
    def test_get_openai_client_initialization_error(self, mock_openai):
        """Test OpenAI client initialization handles errors."""
        # Reset the global client
        import active_interview_app.openai_utils as utils
        utils._openai_client = None

        # Mock initialization failure
        mock_openai.side_effect = Exception("Connection error")

        # Attempt to get client
        with self.assertRaises(ValueError) as context:
            get_openai_client()

        self.assertIn("Failed to initialize OpenAI client",
                      str(context.exception))

    @override_settings(OPENAI_API_KEY="test-key")
    @patch('active_interview_app.openai_utils.OpenAI')
    def testai_available_returns_true(self, mock_openai):
        """Test ai_available returns True when client initializes."""
        # Reset the global client
        import active_interview_app.openai_utils as utils
        utils._openai_client = None

        mock_openai.return_value = MagicMock()

        result = ai_available()
        self.assertTrue(result)

    @override_settings(OPENAI_API_KEY="")
    def testai_available_returns_false(self):
        """Test ai_available returns False when client fails."""
        # Reset the global client
        import active_interview_app.openai_utils as utils
        utils._openai_client = None

        result = ai_available()
        self.assertFalse(result)

    def test_max_tokens_constant(self):
        """Test MAX_TOKENS constant is defined."""
        self.assertEqual(MAX_TOKENS, 15000)


# === Tests for resume_parser.py ===

class ResumeParserTests(TestCase):
    """Test AI-powered resume parsing functions."""

    @override_settings(OPENAI_API_KEY="test-key")
    @patch('active_interview_app.resume_parser.get_client_and_model')
    @patch('active_interview_app.resume_parser.ai_available')
    def test_parse_resume_with_ai_success(
            self, mockai_available, mock_get_client):
        """Test successful resume parsing."""
        # Mock AI availability
        mockai_available.return_value = True

        # Mock OpenAI client response with proper nested structure
        mock_client = MagicMock()
        mock_message = MagicMock()
        mock_message.content = json.dumps(create_sample_parsed_data())
        mock_choice = MagicMock()
        mock_choice.message = mock_message
        mock_response = MagicMock()
        mock_response.choices = [mock_choice]
        mock_client.chat.completions.create.return_value = mock_response
        mock_get_client.return_value = (mock_client, 'gpt-4o', {'tier': 'premium'})

        # Parse resume
        result = parse_resume_with_ai(create_sample_resume_text())

        # Verify structure
        self.assertIn('skills', result)
        self.assertIn('experience', result)
        self.assertIn('education', result)
        self.assertIsInstance(result['skills'], list)
        self.assertIsInstance(result['experience'], list)
        self.assertIsInstance(result['education'], list)

        # Verify content
        self.assertIn("Python", result['skills'])
        self.assertIn("Django", result['skills'])
        self.assertEqual(len(result['experience']), 2)
        self.assertEqual(result['experience'][0]['company'], "Tech Corp")

    @patch('active_interview_app.resume_parser.ai_available')
    def test_parse_resume_ai_unavailable(self, mockai_available):
        """Test parsing fails gracefully when AI is unavailable."""
        mockai_available.return_value = False

        with self.assertRaises(ValueError) as context:
            parse_resume_with_ai(create_sample_resume_text())

        self.assertIn("OpenAI service is unavailable", str(context.exception))

    @override_settings(OPENAI_API_KEY="test-key")
    @patch('active_interview_app.resume_parser.get_client_and_model')
    @patch('active_interview_app.resume_parser.ai_available')
    def test_parse_resume_invalid_json(
            self, mockai_available, mock_get_client):
        """Test parsing handles invalid JSON response."""
        mockai_available.return_value = True

        # Mock client with invalid JSON
        mock_client = MagicMock()
        mock_message = MagicMock()
        mock_message.content = "Not valid JSON"
        mock_choice = MagicMock()
        mock_choice.message = mock_message
        mock_response = MagicMock()
        mock_response.choices = [mock_choice]
        mock_client.chat.completions.create.return_value = mock_response
        mock_get_client.return_value = (mock_client, 'gpt-4o', {'tier': 'premium'})

        with self.assertRaises(ValueError) as context:
            parse_resume_with_ai(create_sample_resume_text())

        self.assertIn("Failed to parse OpenAI response as JSON",
                      str(context.exception))

    @override_settings(OPENAI_API_KEY="test-key")
    @patch('active_interview_app.resume_parser.get_client_and_model')
    @patch('active_interview_app.resume_parser.ai_available')
    def test_parse_resume_with_markdown_json(
            self, mockai_available, mock_get_client):
        """Test parsing handles JSON wrapped in markdown code blocks."""
        mockai_available.return_value = True

        # Mock client with markdown-wrapped JSON
        mock_client = MagicMock()
        markdown_json = "```json\n" + json.dumps(create_sample_parsed_data()) + "\n```"
        mock_message = MagicMock()
        mock_message.content = markdown_json
        mock_choice = MagicMock()
        mock_choice.message = mock_message
        mock_response = MagicMock()
        mock_response.choices = [mock_choice]
        mock_client.chat.completions.create.return_value = mock_response
        mock_get_client.return_value = (mock_client, 'gpt-4o', {'tier': 'premium'})

        # Should successfully parse despite markdown wrapper
        result = parse_resume_with_ai(create_sample_resume_text())

        self.assertIn('skills', result)
        self.assertIsInstance(result['skills'], list)

    @override_settings(OPENAI_API_KEY="test-key")
    @patch('active_interview_app.resume_parser.get_client_and_model')
    @patch('active_interview_app.resume_parser.ai_available')
    def test_parse_resume_missing_fields(
            self, mockai_available, mock_get_client):
        """Test parsing handles missing fields in response."""
        mockai_available.return_value = True

        # Mock client with partial data
        mock_client = MagicMock()
        partial_data = {"skills": ["Python"]}  # Missing experience and education
        mock_message = MagicMock()
        mock_message.content = json.dumps(partial_data)
        mock_choice = MagicMock()
        mock_choice.message = mock_message
        mock_response = MagicMock()
        mock_response.choices = [mock_choice]
        mock_client.chat.completions.create.return_value = mock_response
        mock_get_client.return_value = (mock_client, 'gpt-4o', {'tier': 'premium'})

        result = parse_resume_with_ai(create_sample_resume_text())

        # Should have empty lists for missing fields
        self.assertEqual(result['skills'], ["Python"])
        self.assertEqual(result['experience'], [])
        self.assertEqual(result['education'], [])

    @override_settings(OPENAI_API_KEY="test-key")
    @patch('active_interview_app.resume_parser.get_client_and_model')
    @patch('active_interview_app.resume_parser.ai_available')
    def test_parse_resume_truncates_long_content(
            self, mockai_available, mock_get_client):
        """Test parsing truncates very long resume content."""
        mockai_available.return_value = True

        # Mock successful response
        mock_client = MagicMock()
        mock_message = MagicMock()
        mock_message.content = json.dumps(create_sample_parsed_data())
        mock_choice = MagicMock()
        mock_choice.message = mock_message
        mock_response = MagicMock()
        mock_response.choices = [mock_choice]
        mock_client.chat.completions.create.return_value = mock_response
        mock_get_client.return_value = (mock_client, 'gpt-4o', {'tier': 'premium'})

        # Create very long content (>10,000 chars)
        long_content = "A" * 15000

        result = parse_resume_with_ai(long_content)

        # Should succeed (truncated internally)
        self.assertIn('skills', result)

    def test_validate_parsed_data_valid(self):
        """Test validate_parsed_data with valid structure."""
        data = create_sample_parsed_data()
        self.assertTrue(validate_parsed_data(data))

    def test_validate_parsed_data_missing_keys(self):
        """Test validate_parsed_data with missing keys."""
        invalid_data = {"skills": []}  # Missing experience and education
        self.assertFalse(validate_parsed_data(invalid_data))

    def test_validate_parsed_data_wrong_types(self):
        """Test validate_parsed_data with wrong types."""
        invalid_data = {
            "skills": "Python",  # Should be list, not string
            "experience": [],
            "education": []
        }
        self.assertFalse(validate_parsed_data(invalid_data))

    def test_validate_parsed_data_not_dict(self):
        """Test validate_parsed_data with non-dict input."""
        self.assertFalse(validate_parsed_data("not a dict"))
        self.assertFalse(validate_parsed_data([]))
        self.assertFalse(validate_parsed_data(None))


# === Tests for views.py Integration ===

class ResumeUploadParsingIntegrationTests(TestCase):
    """Test integration of parsing into the upload flow."""

    def setUp(self):
        """Set up test fixtures."""
        self.user = create_test_user()
        self.client.login(username="testuser", password=TEST_PASSWORD)

    @override_settings(OPENAI_API_KEY="test-key")
    @patch('active_interview_app.views.ai_available')
    @patch('active_interview_app.views.parse_resume_with_ai')
    @patch('filetype.guess')
    @patch('pymupdf4llm.to_markdown')
    def test_resume_upload_with_successful_parsing(
        self, mock_to_markdown, mock_guess, mock_parse, mockai_available
    ):
        """Test resume upload triggers parsing successfully."""
        # Mock file type detection
        mock_guess.return_value = type("obj", (object,), {"extension": "pdf"})

        # Mock text extraction
        resume_text = create_sample_resume_text()
        mock_to_markdown.return_value = resume_text

        # Mock AI availability and parsing
        mockai_available.return_value = True
        mock_parse.return_value = create_sample_parsed_data()

        # Upload resume
        file = SimpleUploadedFile(
            "resume.pdf", b"%PDF-1.4 test", content_type="application/pdf")
        response = self.client.post(
            reverse("upload_file"),
            {"file": file, "title": "My Resume"},
            follow=True
        )

        # Verify success message
        self.assertContains(
            response, "Resume uploaded and parsed successfully!")

        # Verify resume was saved with parsed data
        resume = UploadedResume.objects.get(title="My Resume")
        self.assertEqual(resume.parsing_status, 'success')
        self.assertIsNotNone(resume.parsed_at)
        self.assertEqual(len(resume.skills), 6)
        self.assertIn("Python", resume.skills)
        self.assertEqual(len(resume.experience), 2)
        self.assertEqual(resume.experience[0]['company'], "Tech Corp")
        self.assertEqual(len(resume.education), 1)

    @override_settings(OPENAI_API_KEY="test-key")
    @patch('active_interview_app.views.ai_available')
    @patch('active_interview_app.views.parse_resume_with_ai')
    @patch('filetype.guess')
    @patch('pymupdf4llm.to_markdown')
    def test_resume_upload_with_parsing_error(
        self, mock_to_markdown, mock_guess, mock_parse, mockai_available
    ):
        """Test resume upload handles parsing errors gracefully."""
        # Mock file type detection and text extraction
        mock_guess.return_value = type("obj", (object,), {"extension": "pdf"})
        mock_to_markdown.return_value = create_sample_resume_text()

        # Mock AI available but parsing fails
        mockai_available.return_value = True
        mock_parse.side_effect = ValueError("OpenAI rate limit exceeded")

        # Upload resume
        file = SimpleUploadedFile(
            "resume.pdf", b"%PDF-1.4 test", content_type="application/pdf")
        response = self.client.post(
            reverse("upload_file"),
            {"file": file, "title": "My Resume"},
            follow=True
        )

        # Verify warning message
        self.assertContains(response, "Resume uploaded but parsing failed")
        self.assertContains(response, "OpenAI rate limit exceeded")

        # Verify resume was saved with error status
        resume = UploadedResume.objects.get(title="My Resume")
        self.assertEqual(resume.parsing_status, 'error')
        self.assertIsNotNone(resume.parsing_error)
        self.assertIn("OpenAI rate limit exceeded", resume.parsing_error)
        self.assertEqual(resume.skills, [])  # Default empty list

    @override_settings(OPENAI_API_KEY="")
    @patch('active_interview_app.views.ai_available')
    @patch('filetype.guess')
    @patch('pymupdf4llm.to_markdown')
    def test_resume_upload_with_ai_unavailable(
        self, mock_to_markdown, mock_guess, mockai_available
    ):
        """Test resume upload when AI service is unavailable."""
        # Mock file type detection and text extraction
        mock_guess.return_value = type("obj", (object,), {"extension": "pdf"})
        mock_to_markdown.return_value = create_sample_resume_text()

        # Mock AI unavailable
        mockai_available.return_value = False

        # Upload resume
        file = SimpleUploadedFile(
            "resume.pdf", b"%PDF-1.4 test", content_type="application/pdf")
        response = self.client.post(
            reverse("upload_file"),
            {"file": file, "title": "My Resume"},
            follow=True
        )

        # Verify warning message
        self.assertContains(
            response,
            "Resume uploaded but AI parsing is currently unavailable")

        # Verify resume was saved with error status
        resume = UploadedResume.objects.get(title="My Resume")
        self.assertEqual(resume.parsing_status, 'error')
        self.assertEqual(resume.parsing_error, "AI service unavailable")
        self.assertEqual(resume.skills, [])

    @override_settings(OPENAI_API_KEY="test-key")
    @patch('active_interview_app.views.ai_available')
    @patch('filetype.guess')
    def test_docx_upload_with_parsing(self, mock_guess, mockai_available):
        """Test DOCX resume upload and parsing."""
        # Mock AI unavailable to simplify test (focus on file handling)
        mockai_available.return_value = False

        # Create DOCX file
        doc = DocxDocument()
        doc.add_paragraph("John Doe - Software Engineer")
        doc.add_paragraph("Skills: Python, Django")
        fake_file = BytesIO()
        doc.save(fake_file)
        fake_file.seek(0)

        # Mock file type detection
        mock_guess.return_value = type("obj", (object,), {"extension": "docx"})

        # Upload DOCX
        file = SimpleUploadedFile(
            "resume.docx",
            fake_file.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response = self.client.post(
            reverse("upload_file"),
            {"file": file, "title": "My DOCX Resume"},
            follow=True
        )

        # Verify upload succeeded
        self.assertEqual(response.status_code, 200)

        # Verify resume was saved
        resume = UploadedResume.objects.get(title="My DOCX Resume")
        self.assertIn("John Doe", resume.content)

    def test_resume_model_default_parsing_status(self):
        """Test UploadedResume model has correct default parsing_status."""
        resume = UploadedResume.objects.create(
            user=self.user,
            title="Test Resume",
            content="Test content",
            original_filename="test.pdf"
        )

        self.assertEqual(resume.parsing_status, 'pending')
        self.assertEqual(resume.skills, [])
        self.assertEqual(resume.experience, [])
        self.assertEqual(resume.education, [])
        self.assertIsNone(resume.parsing_error)
        self.assertIsNone(resume.parsed_at)
